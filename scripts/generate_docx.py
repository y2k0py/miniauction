#!/usr/bin/env python3
"""Generate simple project documentation as DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "dokumentacja-projektu.docx"
SCREENSHOTS = ROOT / "docs" / "screenshots"


def add_title(doc: Document, text: str) -> None:
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_text(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def build_document() -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    add_title(doc, "MiniAukcje")
    add_text(doc, "Inżynieria Oprogramowania — projekt końcowy")
    add_text(doc, "Python · Django · SQLite · projekt indywidualny")
    doc.add_paragraph()

    # 1. Opis
    add_heading(doc, "1. Opis projektu")
    add_text(
        doc,
        "MiniAukcje to prosta platforma aukcyjna. Użytkownik może wystawić przedmiot, "
        "inni mogą licytować, a po czasie system wybiera zwycięzcę.",
    )
    add_text(doc, "Bez płatności i wysyłki — skupienie na samym procesie aukcji.")

    add_heading(doc, "Funkcje", 2)
    add_bullets(
        doc,
        [
            "Rejestracja i logowanie",
            "Tworzenie aukcji",
            "Licytacja",
            "Kategorie i wyszukiwanie",
            "Lista obserwowanych",
            "Powiadomienia e-mail",
            "Zakończenie aukcji i wybór zwycięzcy",
            "Usuwanie aukcji (bez ofert)",
        ],
    )

    add_heading(doc, "Moduły", 2)
    add_bullets(doc, ["accounts — konta", "auctions — aukcje i licytacja", "notifications — e-mail"])

    # 2. Plan
    add_heading(doc, "2. Plan i zasoby")
    add_table(
        doc,
        ["Tydzień", "Co robione", "Godziny"],
        [
            ["1", "Setup, modele, logowanie", "10 h"],
            ["2", "Aukcje i licytacja", "12 h"],
            ["3", "Watchlist, maile, testy", "10 h"],
            ["4", "Dokumentacja i prezentacja", "13 h"],
        ],
    )
    add_text(doc, "Wykonawca: 1 osoba. Sprzęt: laptop. Baza: SQLite.")

    # 3. Ryzyka i koszty
    add_heading(doc, "3. Ryzyka i koszty")
    add_bullets(
        doc,
        [
            "Opóźnienie → małe etapy pracy",
            "Błąd licytacji → testy automatyczne",
            "Problemy z e-mail → wysyłka w konsoli",
        ],
    )
    add_text(doc, "Koszt szacunkowy: ok. 2500 PLN (50 h × 50 PLN/h). Hosting opcjonalny: 0–25 PLN/mies.")

    # 4. Use cases
    add_heading(doc, "4. Przypadki użycia")
    add_bullets(
        doc,
        [
            "Gość — przegląda aukcje",
            "Użytkownik — licytuje, obserwuje aukcje",
            "Sprzedawca — dodaje i usuwa aukcje",
            "System — kończy aukcje i wysyła maile",
        ],
    )
    add_text(doc, "Logowanie wymagane do: licytacji, dodania aukcji, watchlist, profilu.")

    # 5. Bezpieczeństwo
    add_heading(doc, "5. Bezpieczeństwo")
    add_bullets(
        doc,
        [
            "Hasła hashowane (Django)",
            "CSRF na formularzach",
            "Dostęp tylko dla zalogowanych (licytacja, dodawanie)",
            "Sprzedawca nie licytuje własnej aukcji",
            "Tylko właściciel usuwa swoją aukcję",
        ],
    )

    # 6. Testy
    add_heading(doc, "6. Testy")
    add_text(doc, "18 testów automatycznych — wszystkie przechodzą.")
    add_code(doc, "uv run python manage.py test")
    add_text(doc, "Testowane: logowanie, licytacja, watchlist, usuwanie aukcji, zakończenie aukcji.")

    # 7. Instrukcja
    add_heading(doc, "7. Jak uruchomić")
    add_code(
        doc,
        "uv sync\n"
        "uv run python manage.py migrate\n"
        "uv run python manage.py seed_categories\n"
        "uv run python manage.py seed_demo\n"
        "uv run python manage.py runserver",
    )
    add_text(doc, "Strona: http://127.0.0.1:8000/")
    add_text(doc, "Konta demo: sprzedawca1 / kupujacy1 — hasło: demo12345")

    add_heading(doc, "Jak korzystać", 2)
    add_bullets(
        doc,
        [
            "Zarejestruj się lub zaloguj",
            "Dodaj aukcję (sprzedawca)",
            "Wejdź w aukcję i złóż ofertę (kupujący)",
            "Kliknij Obserwuj, żeby dostać maila o nowej ofercie",
            "Po czasie uruchom: uv run python manage.py end_auctions",
        ],
    )

    # Screenshots
    screenshots = [
        ("01-lista-aukcji.png", "Lista aukcji"),
        ("05-szczegoly-aukcji.png", "Szczegóły aukcji"),
        ("06-profil.png", "Profil"),
    ]
    for filename, caption in screenshots:
        path = SCREENSHOTS / filename
        if path.exists():
            doc.add_picture(str(path), width=Inches(5.0))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 8. Wymagania
    add_heading(doc, "8. Wymagania")
    add_table(
        doc,
        ["Element", "Wersja"],
        [
            ["Python", "3.14+"],
            ["uv", "najnowszy"],
            ["Przeglądarka", "Chrome / Firefox / Safari"],
        ],
    )

    p = doc.add_paragraph("Dokumentacja projektu MiniAukcje — czerwiec 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
